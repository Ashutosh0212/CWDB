


from .forms import CustomUserCreationForm


# class SignUp(generic.CreateView):
#     form_class = CustomUserCreationForm
#     success_url = reverse_lazy("login")
#     template_name = "signup.html"ss

from django.http import HttpResponse  
from django.shortcuts import render, redirect  
from django.contrib.auth import login, authenticate  
from django.contrib.sites.shortcuts import get_current_site  
from django.utils.encoding import force_bytes, force_str  
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode  
from django.template.loader import render_to_string  
from authapp.tokens import account_activation_token  
from django.contrib.auth.models import User  
from django.core.mail import EmailMessage,send_mail
from src.settings import EMAIL_HOST_USER
def signup(request):  
    if request.method == 'POST':  
        form = CustomUserCreationForm(request.POST)  
        if form.is_valid():  
            # save form in the memory not in database  
            user = form.save(commit=False)  
            user.is_active = False  
            user.save()  
            # to get the domain of the current site  
            current_site = get_current_site(request)  
            mail_subject = 'Activation link has been sent to your email id'  
            message = render_to_string('registration/acc_active_email.html', {  
                'user': user,  
                'domain': current_site.domain,  
                'uid':urlsafe_base64_encode(force_bytes(user.pk)),  
                'token':account_activation_token.make_token(user),  
            })  
            to_email = form.cleaned_data.get('email')  
            send_mail(mail_subject, message, EMAIL_HOST_USER, [to_email])
            print(EMAIL_HOST_USER)
            print(to_email)
            print("Email Content:")
            print(message) 
              
            # print(email)
            return HttpResponse('Please confirm your email address to complete the registration')  
    else:  
        form = CustomUserCreationForm()  
    return render(request, 'registration/signup.html', {'form': form})  

from django.contrib.auth.views import LoginView
from django.shortcuts import redirect



from django.contrib.auth import get_user_model

def activate(request, uidb64, token):  
    User = get_user_model()  
    try:  
        uid = force_str(urlsafe_base64_decode(uidb64))  
        user = User.objects.get(pk=uid)  
    except(TypeError, ValueError, OverflowError, User.DoesNotExist):  
        user = None  
    if user is not None and account_activation_token.check_token(user, token):  
        user.is_active = True  
        user.save()  
        return HttpResponse('Thank you for your email confirmation. Now you can login your account.')  
    else:  
        return HttpResponse('Activation link is invalid!')  
    
# views.py


#profile
from django.shortcuts import render, redirect
from django.views import View
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from .models import CustomUser
from .forms import CustomUserChangeForm

@method_decorator(login_required, name='dispatch')
class UserProfileView(View):
    template_name = 'registration/user_profile.html'

    def get(self, request, *args, **kwargs):
        user = request.user
        return render(request, self.template_name, {'user': user})

@method_decorator(login_required, name='dispatch')
class EditProfileView(View):
    template_name = 'registration/edit_profile.html'

    def get(self, request, *args, **kwargs):
        user = request.user
        form = CustomUserChangeForm(instance=user)
        return render(request, self.template_name, {'user': user, 'form': form})

    def post(self, request, *args, **kwargs):
        user = request.user
        form = CustomUserChangeForm(request.POST, instance=user)
        if form.is_valid():
            form.save()
            return redirect('authapp:user_profile')  # Redirect to the user profile page after a successful update
        return render(request, self.template_name, {'user': user, 'form': form})





#admin
# views.py
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.core.mail import send_mail
from django.contrib.auth.decorators import user_passes_test
from .forms import ProposalApprovalForm
from .models import Proposal



from django.core.mail import EmailMessage
from django.template.loader import render_to_string
from django.utils.html import strip_tags
from .models import Notification
from django.utils import timezone

def send_status_change_notification(user_email, project_id, new_status, sanction_letter=None):
    attachment_name = None
    subject = 'Your Proposal Status has Changed'
    context = {'project_id': project_id, 'new_status': new_status}
    message_html = render_to_string('email/notification_template.html', context)
    message_plain = strip_tags(message_html)
    
    notification_message_html=render_to_string('email/notification_template_dashboard.html', context)
    notification_message_plain=strip_tags(notification_message_html)
    # Save the notification to the model
    notification = Notification.objects.create(
        user=user_email,  # Assuming user_email is a User object
        message=notification_message_plain,  # You can customize this based on your needs
        created_at=timezone.now(),
    )
    
    email = EmailMessage(subject, message_plain, from_email=EMAIL_HOST_USER, to=[user_email])
    email.content_subtype = 'html'  # Set the content type to HTML

    if new_status == "Approved" and sanction_letter:
        attachment_name = 'sanction_letter.pdf'
        email.attach(attachment_name, sanction_letter.read(), 'application/pdf')
    
    # Save the attachment to the notification model
   
    notification.save()

    email.send()


#show notifications view
# views.py
from django.shortcuts import render
from .models import Notification
from django.contrib.auth.decorators import login_required

# @login_required
# def show_notifications(request):
#     # Retrieve all notifications for the logged-in user
#     notifications = Notification.objects.filter(user=request.user).order_by('-created_at')
#     first_three_notifications = notifications[:3]

#     # # Mark all new notifications as read
#     # first_three_notifications = notifications.filter(is_read=False)
#     # first_three_notifications.update(is_read=True)
#     # print(first_three_notifications)
#     return render(request, 'main/dashboard.html', {'new_notifications': first_three_notifications})

@login_required
def show_all_notifications(request):
    # Retrieve all notifications for the logged-in user
    all_notifications = Notification.objects.filter(user=request.user).order_by('-created_at')

    return render(request, 'all_notifications.html', {'all_notifications': all_notifications})

from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
# from .models import FormSubmission
from django.utils.crypto import get_random_string

@login_required
def dashboard(request):
    #top three notificatons
    notifications = Notification.objects.filter(user=request.user).order_by('-created_at')
    first_three_notifications = notifications[:3]
     # Count of all proposals submitted by the logged-in user
    submitted_proposals_count = Proposal.objects.filter(user=request.user).count()
    # Count of pending proposals
    pending_proposals_count = Proposal.objects.filter(user=request.user, status='Pending').count()
     # Count of approved proposals
    approved_proposals_count = Proposal.objects.filter(user=request.user, status='Approved').count()
 # Count of proposals to be resubmitted
    resubmit_proposals_count = Proposal.objects.filter(user=request.user, status='Resubmitted').count()
    #count proposals rejected
    rejected_proposals_count=Proposal.objects.filter(user=request.user, status='Rejected').count()

    return render(request, 'main/dashboard.html',{'submitted_proposals_count': submitted_proposals_count,
        'pending_proposals_count': pending_proposals_count,
        'approved_proposals_count': approved_proposals_count,
        'rejected_proposals_count': rejected_proposals_count,'new_notifications': first_three_notifications})

def index(request):
    return render(request,'main/index.html')

from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
# from .forms import ProposalForm
from .models import Proposal
import random
import string

# @login_required
# def submit_proposal(request):
#     if request.method == 'POST':
#         form = ProposalForm(request.POST)
#         if form.is_valid():
#             proposal = form.save(commit=False)
#             proposal.user = request.user
#             proposal.unique_id = ''.join(random.choices(string.digits, k=8))
#             proposal.save()
#             return redirect('proposal_detail', proposal_id=proposal.id)
#     else:
#         form = ProposalForm()
#     return render(request, 'proposal/submit_proposal.html', {'form': form})

# views.py
from django.shortcuts import render, redirect
from .models import Proposal
from django.contrib.auth.decorators import login_required
import uuid

@login_required
@login_required
def submit_proposal(request):
    if request.method == 'POST':
        # Handle text fields
        agency_address = request.POST.get('agencyAddress')
        project_scheme = request.POST.get('projectScheme')
        scheme_component = request.POST.get('schemeComponent')
        applicant_nature = request.POST.get('applicantNature')
        other_nature = request.POST.get('otherNature')
        brief_of_agency = request.POST.get('agencyActivities')
        objectives_of_project = request.POST.get('Objectives')
        brief_of_project = request.POST.get('Brief')
        justification_of_project = request.POST.get('justification')
        methodology_of_project = request.POST.get('Methodology')
        # expected_outcome = request.POST.get('Outcome')
        scenario_change = request.POST.get('change-scenario')
        # beneficiaries = request.POST.get('beneficiaries')
        mode_of_selection = request.POST.get('mode')
        # component_wise_cost = request.POST.get('project Cost')
        # total_duration = request.POST.get('Total duration')
        location_of_project = request.POST.get('projectLocation')
        associated_agency = request.POST.get('associatedAgency')
        bank_details = request.POST.get('bankAccountDetails')
        nodal_officer_details = request.POST.get('nodalOfficerInfo')
        other_info = request.POST.get('otherInfo')

        # Handle file uploads
        expected_outcome_file = request.FILES.get('outcomeFile')
        beneficiaries_file = request.FILES.get('beneficiariesFile')
        component_wise_cost_file = request.FILES.get('projectCostFile')
        total_duration_file = request.FILES.get('durationFile')
        project_report_file = request.FILES.get('projectReport')
        covering_letter_file = request.FILES.get('coveringLetter')

        # Create and save the Proposal object
        proposal = Proposal(
            user=request.user,
            name_and_address=agency_address,
            project_scheme=project_scheme,
            scheme_component=scheme_component,
            nature_of_applicant=applicant_nature,
            other_nature=other_nature,
            brief_of_agency=brief_of_agency,
            objectives_of_project=objectives_of_project,
            brief_of_project=brief_of_project,
            justification_of_project=justification_of_project,
            methodology_of_project=methodology_of_project,
            # expected_outcome=expected_outcome,
            scenario_change=scenario_change,
            # beneficiaries=beneficiaries,
            mode_of_selection=mode_of_selection,
            # component_wise_cost=component_wise_cost,
            # total_duration=total_duration,
            location_of_project=location_of_project,
            associated_agency=associated_agency,
            bank_details=bank_details,
            nodal_officer_details=nodal_officer_details,
            other_info=other_info,
            expected_outcome=expected_outcome_file,
            beneficiaries=beneficiaries_file,
            component_wise_cost=component_wise_cost_file,
            total_duration=total_duration_file,
            project_report=project_report_file,
            covering_letter=covering_letter_file
        )

        # Generate a unique proposal_id
        proposal.unique_id = generate_unique_id()
        proposal.save()

        return redirect('authapp:proposal_status')

    return render(request, 'proposal/submit_proposal.html')

def generate_unique_id():
    # Generate a unique proposal_id using UUID
    unique_id = uuid.uuid4().hex[:8]
    return unique_id

@login_required
def proposal_status(request):
    # Filter proposals for the logged-in user
    user_proposals = Proposal.objects.filter(user=request.user)

    return render(request, 'proposal/proposal_status.html', {'proposals': user_proposals})

@login_required
def proposal_detail(request, proposal_id):
    proposal = Proposal.objects.get(id=proposal_id)
    return render(request, 'proposal/proposal_detail.html', {'proposal': proposal})

@login_required
def proposal_list(request):
    proposals = Proposal.objects.filter(user=request.user)
    return render(request, 'proposal/proposal_list.html', {'proposals': proposals})

@login_required
def admin_proposal_list(request):
    proposals = Proposal.objects.all()
    return render(request, 'proposal/admin_proposal_list.html', {'proposals': proposals})

@login_required
def admin_proposal_detail(request, proposal_id):
    proposal = Proposal.objects.get(id=proposal_id)
    if request.method == 'POST':
        status = request.POST.get('status')
        proposal.status = status
        proposal.save()
        return redirect('admin_proposal_detail', proposal_id=proposal.id)
    return render(request, 'proposal/admin_proposal_detail.html', {'proposal': proposal})


from django.shortcuts import render, redirect
from .models import WMS_RevolvingFund
from .forms import WMSRevolvingFundForm
from django.http import HttpResponse

def get_financial_year():
    # Logic to determine the Indian financial year
    # This is just a sample logic, you can adjust it as per your needs
    from datetime import datetime
    today = datetime.today()
    if today.month >= 4:
        return f'{today.year}-{today.year + 1}'
    else:
        return f'{today.year - 1}-{today.year}'

#generate progress report 
from io import BytesIO
from django.http import HttpResponse
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.shortcuts import render
from django import forms
from .models import ProgressReportDocument, Proposal
from docx import Document
from django.conf import settings
from urllib.parse import quote

def generate_progress_report_document(proposal_unique_id, form_data):
    proposal = Proposal.objects.get(unique_id=proposal_unique_id)
    document = Document()
    document.add_heading(f'Revolving Fund Progress Report - Proposal ID: {proposal}', level=1)

    # # Add form data to the document
    # for field_name, field_value in form_data.items():
    #     document.add_paragraph(f"{field_name}: {field_value}")
    
    # Add form data to the document
    # Add form data to the document
    for field_name, field_value in form_data.items():
        if hasattr(field_value, 'file') and callable(getattr(field_value, 'file', None)):
            document_path = quote(f'media/documents/{field_name}')
            document_url = f'{settings.MEDIA_URL}{document_path}'

        # Add the clickable link to the document
            document.add_paragraph(f"{field_name}: <a href='{document_url}'>{field_name}</a>", style='Hyperlink')
        else:
            document.add_paragraph(f"{field_name}: {field_value}")
    
    

    # Save the document content to a BytesIO object
    temp_file = BytesIO()
    document.save(temp_file)

    # Create a ContentFile from the BytesIO content
    content_file = ContentFile(temp_file.getvalue())

    # Save the document to the database
    report_document = ProgressReportDocument(
        proposal_unique_id=proposal,
        quarter=form_data['quarter'],
        financial_year=form_data['financial_year'],
    )
    report_document.document.save(f'{proposal_unique_id}_progress_report.docx', content_file)
    report_document.save()

    # Clean up the BytesIO object
    temp_file.close()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=revolving_fund_progress_report_{proposal_unique_id}.docx'
    document.save(response)

    return response




def revolving_fund_progress_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = WMSRevolvingFundForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            # Generate and save progress report document
            generate_progress_report_document(proposal_unique_id, form.cleaned_data)
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WMSRevolvingFundForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'progressReports/WMS/1.RevolvingFund.html', context)

from .models import EPortal
from .forms import EPortalForm
from django.http import HttpResponse

from django.shortcuts import render
from django.http import HttpResponse
from .forms import EPortalForm
from .models import EPortal

def eportal_progress_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = EPortalForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()
            form_instance.save()
            generate_progress_report_document(proposal_unique_id, form.cleaned_data)
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),
        }
        form = EPortalForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'progressReports/WMS/2.EPortal.html', context)

@login_required
def progress_report(request):
    user = request.user
    proposals = Proposal.objects.filter(user=user, status='Approved')
    return render(request, 'progressReports/progressreport.html', {'proposals': proposals})

from django.shortcuts import render, redirect
from .models import WMS_SelfHelpGroup
from .forms import WMSSelfHelpGroupForm
from django.http import HttpResponse

def selfhelpgroup_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = WMSSelfHelpGroupForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WMSSelfHelpGroupForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'progressReports/WMS/3.SelfHelpGroup.html', context)

from django.shortcuts import render, redirect
from .models import WMS_BuyerSellerExpo
from .forms import WMSBuyerSellerExpoForm
from django.http import HttpResponse

def buyersellerexpo_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = WMSBuyerSellerExpoForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WMSBuyerSellerExpoForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'progressReports/WMS/4.BuyerSellerExpo.html', context)

from django.shortcuts import render, redirect
from .models import WMS_InfrastructureDevelopment
from .forms import WMSInfrastructureDevelopmentForm
from django.http import HttpResponse

def infrastructuredevelopment_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = WMSInfrastructureDevelopmentForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WMSInfrastructureDevelopmentForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'progressReports/WMS/5.infrastructureDevelopment.html', context)

from .models import WoolenExpo, WoolenExpoHiring
from .forms import WoolenExpoForm, WoolenExpoHiringForm

def woolen_expo(request,proposal_unique_id):
    if request.method == 'POST':
        form = WoolenExpoForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WoolenExpoForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'progressReports/WMS/6.WoolenExpo.html', context)
    
def woolen_expo_hiring(request,proposal_unique_id):
    if request.method == 'POST':
        form = WoolenExpoHiringForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WoolenExpoHiringForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'progressReports/WMS/7.WoolenExpoHiring.html', context)

# views.py
from django.shortcuts import render, redirect
from .models import WPS_CFC
from .forms import WPSCFCForm
from django.http import HttpResponse

def cfc_progress_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = WPSCFCForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WPSCFCForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'progressReports/WPS/1.CFC.html', context)

from .forms import WPSSheepShearingMachingForm
def sheep_shearing_machining(request, proposal_unique_id):
    if request.method == 'POST':
        form = WPSSheepShearingMachingForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WPSSheepShearingMachingForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'progressReports/WPS/2.SheepShearingMaching.html', context)

from .forms import WPSEquipmentForm
def equipment(request, proposal_unique_id):
    if request.method == 'POST':
        form = WPSEquipmentForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WPSEquipmentForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'progressReports/WPS/3.Equipment.html', context)

from django.shortcuts import render
from django.http import HttpResponse
from .forms import WPSSmallToolsDistributionForm

def small_tools_distribution(request, proposal_unique_id):
    if request.method == 'POST':
        form = WPSSmallToolsDistributionForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WPSSmallToolsDistributionForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'progressReports/WPS/4.SmallToolsDistribution.html', context)

from .models import HRD_ShortTermProgramme
from .forms import HRDShortTermProgrammeForm

def short_term_programme(request, proposal_unique_id):
    if request.method == 'POST':
        form = HRDShortTermProgrammeForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),
        }
        form = HRDShortTermProgrammeForm(initial=initial_data)

    context = {'form': form}
    return render(request, 'progressReports/HRD/1.ShortTermProgramme.html', context)

from .models import HRD_OnsiteTraining
from .forms import HRDOnsiteTrainingForm

def onsite_training_progress_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = HRDOnsiteTrainingForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = HRDOnsiteTrainingForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'progressReports/HRD/2.OnsiteTraining.html', context)

from django.shortcuts import render, redirect
from .models import HRD_ShearingMachineTraining
from .forms import HRDShearingMachineTrainingForm
from django.http import HttpResponse

def shearing_machine_training_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = HRDShearingMachineTrainingForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = HRDShearingMachineTrainingForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'progressReports/HRD/3.ShearingMachineTraining.html', context)

from .models import RD
from .forms import RDForm

def rd_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = RDForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = RDForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'progressReports/HRD/4.R&D.html', context)

from django.shortcuts import render, redirect
from .models import DomesticMeeting
from .forms import DomesticMeetingForm
from django.http import HttpResponse

def domestic_meeting_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = DomesticMeetingForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = DomesticMeetingForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'progressReports/HRD/5.DomesticMeeting.html', context)

from django.shortcuts import render, redirect
from .models import OrganisingSeminar
from .forms import OrganisingSeminarForm
from django.http import HttpResponse

def organising_seminar_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = OrganisingSeminarForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = OrganisingSeminarForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'progressReports/HRD/6.OrganisingSeminar.html', context)

from django.shortcuts import render, redirect
from .models import WoolSurvey
from .forms import WoolSurveyForm
from django.http import HttpResponse

def wool_survey_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = WoolSurveyForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WoolSurveyForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'progressReports/HRD/7.WoolSurvey.html', context)

from django.shortcuts import render, redirect
from .models import WoolTestingLab
from .forms import WoolTestingLabForm
from django.http import HttpResponse

def wool_testing_lab_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = WoolTestingLabForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = WoolTestingLabForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'progressReports/HRD/8.WoolTestingLab.html', context)

from django.shortcuts import render, redirect
from .models import PublicityMonitoring
from .forms import PublicityMonitoringForm
from django.http import HttpResponse

def publicity_monitoring_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = PublicityMonitoringForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = PublicityMonitoringForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'progressReports/HRD/9.PublicityMonitoring.html', context)

# pwds
from django.shortcuts import render, redirect
from .models import PWDS_PashminaRevolvingFund
from .forms import PWDS_PashminaRevolvingFundForm
from django.http import HttpResponse


def pashmina_revolving_fund_progress_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = PWDS_PashminaRevolvingFundForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Automatically fill the financial year
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = PWDS_PashminaRevolvingFundForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'progressReports/PWDS/1.PashminaRevolvingFund.html', context)

from django.shortcuts import render, redirect
from .models import PWDS_PashminaCFC
from .forms import PWDS_PashminaCFCForm
from django.http import HttpResponse

def pashmina_cfc_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = PWDS_PashminaCFCForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()  # Assuming get_financial_year is defined elsewhere
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),  # Set the initial value for financial_year
        }
        form = PWDS_PashminaCFCForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'PWDS/2.PashminaCFC.html', context)

from django.shortcuts import render, redirect
from .models import ShelterShedConstruction
from .forms import ShelterShedConstructionForm
from django.http import HttpResponse

def shelter_shed_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = ShelterShedConstructionForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),
        }
        form = ShelterShedConstructionForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'ShelterShedConstruction/3.ShelterShedConstruction.html', context)

from django.shortcuts import render, redirect
from .models import PortableTentDist
from .forms import PortableTentDistForm
from django.http import HttpResponse

def portable_tent_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = PortableTentDistForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),
        }
        form = PortableTentDistForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'PortableTentDist/4.PortableTentDist.html', context)

from django.shortcuts import render, redirect
from .models import PredatorProofLightsDist
from .forms import PredatorProofLightsDistForm
from django.http import HttpResponse

def predator_proof_lights_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = PredatorProofLightsDistForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),
        }
        form = PredatorProofLightsDistForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'PredatorProofLightsDist/5.PredatorProofLightsDist.html', context)

from django.shortcuts import render, redirect
from .models import TestingEquipment
from .forms import TestingEquipmentForm
from django.http import HttpResponse

def testing_equipment_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = TestingEquipmentForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),
        }
        form = TestingEquipmentForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'TestingEquipment/6.TestingEquipment.html', context)

from django.shortcuts import render, redirect
from .models import ShowroomDevelopment
from .forms import ShowroomDevelopmentForm
from django.http import HttpResponse

def showroom_development_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = ShowroomDevelopmentForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),
        }
        form = ShowroomDevelopmentForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'ShowroomDevelopment/7.ShowroomDevelopment.html', context)

from django.shortcuts import render, redirect
from .models import FodderLandDevelopment
from .forms import FodderLandDevelopmentForm
from django.http import HttpResponse

def fodder_land_development_report(request, proposal_unique_id):
    if request.method == 'POST':
        form = FodderLandDevelopmentForm(request.POST, request.FILES)
        if form.is_valid():
            form_instance = form.save(commit=False)
            form_instance.financial_year = get_financial_year()
            form_instance.save()
            return HttpResponse('Form submitted successfully!')
    else:
        initial_data = {
            'proposal_unique_id': proposal_unique_id,
            'financial_year': get_financial_year(),
        }
        form = FodderLandDevelopmentForm(initial=initial_data)
    
    context = {'form': form}
    return render(request, 'FodderLandDevelopment/8.FodderLandDevelopment.html', context)


#staff view
from django.contrib.auth.views import LoginView
from django.urls import reverse_lazy
class CustomLoginView(LoginView):
    def get_success_url(self):
        user = self.request.user

        if user.is_superuser or user.is_staff:
            return reverse_lazy('admin:index')  # Use reverse_lazy to avoid URL resolution issues
        return reverse_lazy('authapp:dashboard') 
    
@user_passes_test(lambda u: u.is_staff or u.is_superuser, login_url='/login/')
def staff_dashboard(request):
    # Count the number of proposals for each status
    pending_count = Proposal.objects.filter(status='Pending').count()
    approved_count = Proposal.objects.filter(status='Approved').count()
    completed_count = Proposal.objects.filter(status='Completed').count()
    rejected_count = Proposal.objects.filter(status='Rejected').count()

    # Your view logic goes here

    return render(request, 'staff_template/staff_dashboard.html', {
        'pending_count': pending_count,
        'approved_count': approved_count,
        'completed_count': completed_count,
        'rejected_count': rejected_count,
    })


@user_passes_test(lambda u: u.is_staff or u.is_superuser)
def all_proposal_details(request):
    proposals = Proposal.objects.all()
    return render(request, 'staff_template/admin_proposal_list.html', {'proposals': proposals})

@user_passes_test(lambda u: u.is_staff or u.is_superuser)
def submit_approval(request, proposal_id):
    proposal = get_object_or_404(Proposal, unique_id=proposal_id)

    if request.method == 'POST':
        form = ProposalApprovalForm(request.POST, request.FILES,instance=proposal)
        if form.is_valid():
            # Save the form to update the model
            form.save()

            # Notify the user about the status change
            send_status_change_notification(proposal.user, proposal.unique_id, proposal.status,proposal.sanction_letter)

            messages.success(request, 'Proposal status changed successfully.')
            # return redirect('admin:authapp_proposal_changelist')  # Replace 'proposal_list' with your actual URL name
            return redirect('authapp:all_proposal_details')
    else:
        form = ProposalApprovalForm(instance=proposal)

    return render(request, 'admin/submit_approval.html', {'form': form, 'proposal': proposal})